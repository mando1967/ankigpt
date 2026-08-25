# Copyright: Ankitects Pty Ltd and contributors
# License: GNU AGPL, version 3 or later; http://www.gnu.org/licenses/agpl.html

"""Document reading and the concept extraction pipeline.

Pure Python: runs on a background thread, no Qt or collection access.

Large documents are read *agentically* under a hard per-file character
budget:

1. structure pass (free): split into sections at headings and build a skim
   (title, position, length, first few hundred characters of each section);
2. plan call: the model picks which sections to read, with priorities, given
   the learner's instructions and the budget;
3. extraction over the chosen sections (map stage);
4. one bounded gap call: the model may ask for a few unread sections that
   look important given what was extracted so far;
5. merge/rank (reduce stage).

Small documents skip 1-4 and are read whole. If the plan call fails, evenly
spaced sampling across the document is used as a fallback.
"""

from __future__ import annotations

import math
import os
import re
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from typing import Protocol

from aqt.ankigpt import prompts
from aqt.ankigpt.prompts import ConceptCandidate

SUPPORTED_EXTENSIONS = (".pdf", ".txt", ".md", ".markdown", ".docx")
CHUNK_CHARS = 12_000
CHUNK_OVERLAP = 500
MERGE_BATCH = 60
DEFAULT_MAX_CHARS_PER_FILE = 150_000  # ~37k tokens, ~13 requests
MAX_PER_CHUNK = 15

SEGMENT_SEPARATOR = "\n\n[...]\n\n"
OUTLINE_MAX_CHARS = 3_000
SKIM_MAX_CHARS = 16_000
PREVIEW_MAX_CHARS = 300
MIN_SECTION_CHARS = 1_500
MAX_SECTIONS = 150
MIN_READ_CHARS = 600
GAP_MAX_SECTIONS = 4


class ExtractionError(Exception):
    pass


class Cancelled(Exception):
    pass


class JsonClient(Protocol):
    def complete_json(
        self, system: str, user: str, schema_name: str, json_schema: dict
    ) -> dict: ...


@dataclass
class ReadReport:
    """How much of a document was actually sent to the model."""

    total_chars: int
    chars_read: int
    sections_total: int
    sections_read: int
    planned: bool = False  # model chose the sections
    fallback: bool = False  # even sampling was used instead of a plan

    @property
    def partial(self) -> bool:
        return self.chars_read < self.total_chars

    @property
    def coverage(self) -> float:
        return min(1.0, self.chars_read / self.total_chars) if self.total_chars else 1.0


@dataclass
class Document:
    name: str
    text: str
    total_chars: int = 0
    outline: str = ""
    report: ReadReport | None = None
    path: str = ""
    pages: list[int] = field(default_factory=list)  # char offset of each page start

    def __post_init__(self) -> None:
        if not self.total_chars:
            self.total_chars = len(self.text)
        if not self.outline:
            self.outline = build_outline(self.text)


@dataclass
class Section:
    index: int
    title: str
    start: int
    end: int
    text: str

    @property
    def length(self) -> int:
        return len(self.text)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def is_supported(path: str) -> bool:
    return path.lower().endswith(SUPPORTED_EXTENSIONS)


def extract_text(path: str) -> Document:
    """Read a document to plain text (all of it; budgeting happens later)."""
    name = os.path.basename(path)
    lower = path.lower()
    pages: list[int] = []
    if lower.endswith(".pdf"):
        text, pages = _pdf_text(path)
    elif lower.endswith(".docx"):
        text = _normalize(_docx_text(path))
    elif lower.endswith((".txt", ".md", ".markdown")):
        with open(path, encoding="utf-8", errors="replace") as f:
            text = _normalize(f.read())
    else:
        raise ExtractionError(f"unsupported file type: {name}")
    return Document(name=name, text=text, path=os.path.abspath(path), pages=pages)


def _pdf_text(path: str) -> tuple[str, list[int]]:
    """Normalized text of a PDF plus the offset at which each page starts."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency present in builds
        raise ExtractionError("PDF support requires the pypdf package") from exc
    reader = PdfReader(path)
    parts: list[str] = []
    starts: list[int] = []
    pos = 0
    for page in reader.pages:
        page_text = _normalize(page.extract_text() or "")
        starts.append(pos)
        parts.append(page_text)
        pos += len(page_text) + 2
    return "\n\n".join(parts), starts


def _docx_text(path: str) -> str:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support requires the python-docx package") from exc
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Structure: headings, outline, sections, skim
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(
    r"^(?:#{1,6}\s+\S.*|(?:chapter|part|section|unit|lecture|week|module)\s+[\dIVXivx]+\b.*"
    r"|\d+(?:\.\d+)*\.?\s+[A-Z].*|[A-Z][A-Z0-9 ,:'&-]{3,})$",
    re.IGNORECASE,
)


def _looks_like_heading(line: str) -> bool:
    line = line.strip()
    if not 3 <= len(line) <= 90 or line.endswith((".", ",", ";", ":")):
        return False
    if line.startswith("#"):
        return True
    m = _HEADING_RE.match(line)
    if not m:
        return False
    # all-caps branch must really be mostly letters
    if line.isupper():
        letters = sum(c.isalpha() for c in line)
        return letters >= len(line) * 0.6
    return True


def _heading_title(raw: str) -> str:
    return raw.strip().lstrip("#").strip()


def _headings(text: str) -> list[tuple[int, str]]:
    """(offset, title) for every heading-like line, in document order."""
    out: list[tuple[int, str]] = []
    pos = 0
    for raw in text.split("\n"):
        if _looks_like_heading(raw):
            out.append((pos, _heading_title(raw)))
        pos += len(raw) + 1
    return out


def build_outline(text: str, max_chars: int = OUTLINE_MAX_CHARS) -> str:
    """Cheap, LLM-free table of contents: heading-like lines from the whole
    document with their approximate position, capped at max_chars."""
    if not text:
        return ""
    lines: list[str] = []
    seen: set[str] = set()
    total = len(text)
    for pos, title in _headings(text):
        if title.lower() in seen:
            continue
        seen.add(title.lower())
        lines.append(f"[{int(100 * pos / total):3d}%] {title}")
    out = "\n".join(lines)
    if len(out) > max_chars:
        out = out[:max_chars].rsplit("\n", 1)[0]
    return out


def split_sections(
    text: str,
    min_chars: int = MIN_SECTION_CHARS,
    max_sections: int = MAX_SECTIONS,
) -> list[Section]:
    """Split a document into sections at headings.

    Tiny sections are merged into their predecessor; documents without
    usable headings get fixed-size pseudo-sections so planning still works.
    """
    if not text:
        return []
    bounds = [(pos, title) for pos, title in _headings(text)]
    if len(bounds) < 2:
        bounds = [
            (start, f"Part {i + 1}")
            for i, start in enumerate(range(0, len(text), CHUNK_CHARS))
        ]
    if bounds[0][0] != 0:
        bounds.insert(0, (0, "Preamble"))

    raw: list[tuple[str, int, int]] = []
    for i, (start, title) in enumerate(bounds):
        end = bounds[i + 1][0] if i + 1 < len(bounds) else len(text)
        raw.append((title, start, end))

    # group small sections: a section is absorbed into the previous group
    # while either of them is under min_chars, but groups stop growing at
    # 4 * min_chars so a run of short sections never collapses into one.
    merged: list[tuple[str, int, int]] = []
    for title, start, end in raw:
        cur_len = end - start
        if merged:
            ptitle, pstart, pend = merged[-1]
            prev_len = pend - pstart
            if (
                prev_len < min_chars or cur_len < min_chars
            ) and prev_len < 4 * min_chars:
                keep = ptitle if prev_len >= cur_len else title
                merged[-1] = (keep, pstart, end)
                continue
        merged.append((title, start, end))

    # cap the number of sections by grouping neighbours
    if len(merged) > max_sections:
        group = math.ceil(len(merged) / max_sections)
        grouped: list[tuple[str, int, int]] = []
        for i in range(0, len(merged), group):
            block = merged[i : i + group]
            title = (
                block[0][0] if len(block) == 1 else f"{block[0][0]} … {block[-1][0]}"
            )
            grouped.append((title, block[0][1], block[-1][2]))
        merged = grouped

    return [
        Section(
            index=i, title=title, start=start, end=end, text=text[start:end].strip()
        )
        for i, (title, start, end) in enumerate(merged)
    ]


def _preview(section: Section, max_chars: int) -> str:
    body = section.text
    first_line_end = body.find("\n")
    if first_line_end != -1 and _looks_like_heading(body[:first_line_end]):
        body = body[first_line_end + 1 :]
    body = re.sub(r"\s+", " ", body).strip()
    if len(body) > max_chars:
        body = body[:max_chars].rsplit(" ", 1)[0] + "…"
    return body


def build_skim(
    sections: list[Section], total_chars: int, max_chars: int = SKIM_MAX_CHARS
) -> str:
    """One line per section: index, position, length and a short preview."""
    if not sections:
        return ""
    per_section = max(80, min(PREVIEW_MAX_CHARS, max_chars // len(sections) - 70))
    lines = []
    for s in sections:
        pct = int(100 * s.start / total_chars) if total_chars else 0
        lines.append(
            f"[{s.index}] ({pct}%, {s.length:,} chars) {s.title}: "
            f"{_preview(s, per_section)}"
        )
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Budgeted reading
# ---------------------------------------------------------------------------


def sample_evenly(text: str, budget: int, window: int = CHUNK_CHARS) -> list[str]:
    """Pick evenly spaced windows (on paragraph boundaries) that together
    stay within `budget` characters and span the whole text."""
    if budget <= 0 or not text:
        return []
    if len(text) <= budget:
        return [text]
    # aim for at least four windows so the sample spans the document
    window = max(1000, min(window, budget // 4))
    count = max(1, budget // window)
    window = budget // count
    starts = [
        int(i * (len(text) - window) / max(count - 1, 1)) if count > 1 else 0
        for i in range(count)
    ]
    segments: list[str] = []
    last_end = 0
    for start in starts:
        start = max(start, last_end)
        if start >= len(text):
            break
        # snap forward to a paragraph boundary when one is near
        boundary = text.find("\n\n", start, start + window // 2)
        if boundary != -1 and start != 0:
            start = boundary + 2
        end = min(len(text), start + window)
        if end < len(text):
            cut = text.rfind("\n\n", start + window // 2, end)
            if cut != -1:
                end = cut
        segment = text[start:end].strip()
        if segment:
            segments.append(segment)
        last_end = end
    return segments


@dataclass
class ReadPlan:
    """Sections to read (document order) with how many characters of each."""

    picks: list[tuple[Section, int]] = field(default_factory=list)

    @property
    def chars(self) -> int:
        return sum(take for _, take in self.picks)

    def indices(self) -> set[int]:
        return {s.index for s, _ in self.picks}


def allocate_budget(
    sections: list[Section],
    priorities: list[tuple[int, int]],
    budget: int,
    min_read: int = MIN_READ_CHARS,
) -> ReadPlan:
    """Turn (index, priority) choices into a plan within `budget` chars.

    Higher priority first; ties in document order. A section that does not
    fit entirely is read partially (evenly sampled) if enough room remains.
    """
    by_index = {s.index: s for s in sections}
    seen: set[int] = set()
    ordered: list[tuple[int, Section]] = []
    for index, priority in priorities:
        section = by_index.get(index)
        if section is None or index in seen:
            continue
        seen.add(index)
        ordered.append((priority, section))
    ordered.sort(key=lambda item: (-item[0], item[1].index))

    remaining = budget
    chosen: list[tuple[Section, int]] = []
    for _priority, section in ordered:
        if remaining < min_read:
            break
        take = min(section.length, remaining)
        if take < min_read:
            continue
        chosen.append((section, take))
        remaining -= take
    chosen.sort(key=lambda item: item[0].index)
    return ReadPlan(chosen)


def render_plan(plan: ReadPlan) -> str:
    """The text to send for extraction: chosen sections, partial ones sampled."""
    parts: list[str] = []
    for section, take in plan.picks:
        if take >= section.length:
            parts.append(section.text)
        else:
            parts.extend(sample_evenly(section.text, take))
    return SEGMENT_SEPARATOR.join(p for p in parts if p)


def _plan_with_model(
    doc: Document,
    sections: list[Section],
    instructions: str,
    budget: int,
    client: JsonClient,
) -> ReadPlan:
    skim = build_skim(sections, doc.total_chars)
    system, user = prompts.build_plan_prompt(
        skim, instructions, doc.name, doc.total_chars, budget, len(sections)
    )
    data = client.complete_json(system, user, "plan_reading", prompts.PLAN_SCHEMA)
    priorities = prompts.parse_plan(data)
    return allocate_budget(sections, priorities, budget)


def _gap_sections(
    doc: Document,
    unread: list[Section],
    extracted_titles: list[str],
    instructions: str,
    remaining: int,
    client: JsonClient,
) -> list[int]:
    if not unread or remaining < MIN_READ_CHARS:
        return []
    skim = build_skim(unread, doc.total_chars, max_chars=SKIM_MAX_CHARS // 2)
    system, user = prompts.build_gap_prompt(
        skim, extracted_titles, instructions, doc.name, remaining, GAP_MAX_SECTIONS
    )
    data = client.complete_json(system, user, "find_gaps", prompts.GAP_SCHEMA)
    valid = {s.index for s in unread}
    return [i for i in prompts.parse_gaps(data) if i in valid][:GAP_MAX_SECTIONS]


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def chunk_text(
    text: str, chunk_chars: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """Split on paragraph boundaries into chunks of roughly chunk_chars."""
    if not text:
        return []
    if len(text) <= chunk_chars:
        return [text]
    paragraphs = re.split(r"\n\s*\n", text)
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        # a single paragraph longer than a chunk is split hard
        while len(para) > chunk_chars:
            if current:
                chunks.append("\n\n".join(current))
                current, current_len = [], 0
            chunks.append(para[:chunk_chars])
            para = para[chunk_chars - overlap :]
        if current_len + len(para) + 2 > chunk_chars and current:
            chunks.append("\n\n".join(current))
            # carry over the tail of the previous chunk as overlap
            tail = chunks[-1][-overlap:] if overlap else ""
            current, current_len = ([tail] if tail else []), len(tail)
        current.append(para)
        current_len += len(para) + 2
    if current:
        chunks.append("\n\n".join(current))
    return [c for c in chunks if c.strip()]


def suggest_target_count(total_chars: int) -> int:
    return max(5, min(60, total_chars // 4000))


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------


@dataclass
class ProgressEvent:
    """A progress update from the extraction pipeline.

    stage: plan | extract | gap | merge | done
    current/total: position within the stage
    message: a human-readable line worth showing in a log
    candidates: concept candidates found so far
    """

    stage: str
    current: int
    total: int
    message: str = ""
    candidates: int = 0


ProgressFn = Callable[[ProgressEvent], None]
ReportFn = Callable[[str, int, int, str], None]


@dataclass
class _DocState:
    doc: Document
    sections: list[Section]
    read_indices: set[int]
    budget_left: int
    planned: bool
    fallback: bool
    chars_read: int
    text: str = ""


def extract_concepts(
    docs: list[Document],
    instructions: str,
    target: int,
    client: JsonClient,
    progress: ProgressFn | None = None,
    should_cancel: Callable[[], bool] | None = None,
    workers: int = 4,
    max_chars_per_file: int | None = None,
) -> list[ConceptCandidate]:
    """Extract concepts from documents; `max_chars_per_file` is a hard cap on
    the characters of each document sent to the model (skim included)."""

    found = 0

    def report(stage: str, i: int, n: int, message: str = "") -> None:
        if progress:
            progress(ProgressEvent(stage, i, n, message, found))

    def on_found(count: int) -> None:
        nonlocal found
        found = count

    def check_cancel() -> None:
        if should_cancel and should_cancel():
            raise Cancelled()

    # ---- pass 1: decide what to read from each document
    states: list[_DocState] = []
    texts: list[tuple[Document, str]] = []
    report("plan", 0, len(docs))
    for n, doc in enumerate(docs, start=1):
        check_cancel()
        state = _read_document(doc, instructions, max_chars_per_file, client)
        states.append(state)
        if doc.report is not None and doc.report.chars_read:
            texts.append((doc, _text_for(state)))
        report("plan", n, len(docs), _describe_read(doc))

    # ---- pass 2: extraction over what was read
    candidates = _extract_from_texts(
        texts, instructions, target, client, workers, report, check_cancel, on_found
    )

    # ---- pass 3: one bounded gap check per partially read document
    gap_texts: list[tuple[Document, str]] = []
    partial = [s for s in states if s.planned and s.budget_left >= MIN_READ_CHARS]
    if partial and candidates:
        report("gap", 0, len(partial))
        titles = [c.title for c in candidates]
        for n, state in enumerate(partial, start=1):
            check_cancel()
            unread = [s for s in state.sections if s.index not in state.read_indices]
            try:
                wanted = _gap_sections(
                    state.doc, unread, titles, instructions, state.budget_left, client
                )
            except Exception:
                wanted = []
            message = f"{state.doc.name}: no skipped sections worth reading"
            if wanted:
                plan = allocate_budget(
                    state.sections, [(i, 1) for i in wanted], state.budget_left
                )
                if plan.picks:
                    state.read_indices |= plan.indices()
                    state.budget_left -= plan.chars
                    state.chars_read += plan.chars
                    _update_report(state)
                    gap_texts.append((state.doc, render_plan(plan)))
                    names = ", ".join(s.title for s, _ in plan.picks)
                    message = f"{state.doc.name}: also reading {names}"
            report("gap", n, len(partial), message)
        if gap_texts:
            candidates.extend(
                _extract_from_texts(
                    gap_texts,
                    instructions,
                    target,
                    client,
                    workers,
                    report,
                    check_cancel,
                    on_found,
                    offset=len(candidates),
                )
            )

    if not candidates:
        return []
    report(
        "merge", 0, 1, f"merging {len(candidates)} candidates into {target} concepts"
    )
    merged = _reduce(candidates, instructions, target, client, check_cancel)
    report("merge", 1, 1, f"{len(merged)} concepts ready")
    report("done", 1, 1)
    return merged


def _describe_read(doc: Document) -> str:
    r = doc.report
    if r is None:
        return doc.name
    if not r.partial:
        return f"{doc.name}: {r.total_chars:,} characters, read in full"
    how = "sampled evenly" if r.fallback else "planned"
    return (
        f"{doc.name}: {r.total_chars:,} characters in {r.sections_total} sections; "
        f"reading {r.sections_read} ({int(100 * r.coverage)}%, {how})"
    )


def _read_document(
    doc: Document,
    instructions: str,
    max_chars: int | None,
    client: JsonClient,
) -> _DocState:
    """Choose what to read from `doc` within the budget and fill doc.report."""
    sections = split_sections(doc.text)
    total = doc.total_chars
    if max_chars is None or total <= max_chars:
        doc.report = ReadReport(total, total, len(sections), len(sections))
        return _DocState(
            doc, sections, {s.index for s in sections}, 0, False, False, total, doc.text
        )

    # the skim and outline are sent too; they come out of the same budget
    skim_cost = min(SKIM_MAX_CHARS, len(build_skim(sections, total))) + len(doc.outline)
    budget = max(0, max_chars - skim_cost)
    planned = fallback = False
    try:
        plan = _plan_with_model(doc, sections, instructions, budget, client)
        planned = bool(plan.picks)
    except Exception:
        plan = ReadPlan()
    if not plan.picks:
        # no usable plan: fall back to even sampling across the document
        fallback = True
        segments = sample_evenly(doc.text, budget)
        text = SEGMENT_SEPARATOR.join(segments)
        doc.report = ReadReport(
            total, len(text), len(sections), len(segments), planned=False, fallback=True
        )
        return _DocState(doc, sections, set(), 0, False, True, len(text), text)

    state = _DocState(
        doc,
        sections,
        plan.indices(),
        budget - plan.chars,
        planned,
        fallback,
        plan.chars,
        render_plan(plan),
    )
    _update_report(state)
    return state


def _update_report(state: _DocState) -> None:
    state.doc.report = ReadReport(
        total_chars=state.doc.total_chars,
        chars_read=min(state.chars_read, state.doc.total_chars),
        sections_total=len(state.sections),
        sections_read=len(state.read_indices),
        planned=state.planned,
        fallback=state.fallback,
    )


def _text_for(state: _DocState) -> str:
    return state.text


def _extract_from_texts(
    texts: list[tuple[Document, str]],
    instructions: str,
    target: int,
    client: JsonClient,
    workers: int,
    report: ReportFn,
    check_cancel: Callable[[], None],
    on_found: Callable[[int], None] | None = None,
    offset: int = 0,
) -> list[ConceptCandidate]:
    chunks: list[tuple[Document, str]] = []
    for doc, text in texts:
        for chunk in chunk_text(text):
            chunks.append((doc, chunk))
    if not chunks:
        return []
    total_chars = sum(len(c) for _, c in chunks)

    def want_for(chunk: str) -> int:
        share = target * 1.5 * len(chunk) / max(total_chars, 1)
        return max(2, min(MAX_PER_CHUNK, math.ceil(share) + 2))

    def run_chunk(item: tuple[Document, str]) -> list[ConceptCandidate]:
        check_cancel()
        doc, chunk = item
        partial = doc.report is not None and doc.report.partial
        system, user = prompts.build_extract_prompt(
            chunk,
            instructions,
            want_for(chunk),
            doc.name,
            outline=doc.outline,
            sampled=partial,
        )
        data = client.complete_json(
            system, user, "extract_concepts", prompts.EXTRACT_SCHEMA
        )
        return prompts.parse_concepts(data)

    candidates: list[ConceptCandidate] = []
    report("extract", 0, len(chunks), f"{len(chunks)} chunks to read")
    with ThreadPoolExecutor(max_workers=max(1, workers)) as pool:
        for i, result in enumerate(pool.map(run_chunk, chunks), start=1):
            candidates.extend(result)
            if on_found:
                on_found(offset + len(candidates))
            doc, _chunk = chunks[i - 1]
            report(
                "extract",
                i,
                len(chunks),
                f"{doc.name}: chunk {i}/{len(chunks)} gave {len(result)} candidates "
                f"({offset + len(candidates)} so far)",
            )
            check_cancel()
    return candidates


def _merge_once(
    candidates: list[ConceptCandidate],
    instructions: str,
    target: int,
    client: JsonClient,
) -> list[ConceptCandidate]:
    system, user = prompts.build_merge_prompt(candidates, instructions, target)
    data = client.complete_json(system, user, "merge_concepts", prompts.MERGE_SCHEMA)
    return prompts.parse_concepts(data)


def _reduce(
    candidates: list[ConceptCandidate],
    instructions: str,
    target: int,
    client: JsonClient,
    check_cancel: Callable[[], None],
) -> list[ConceptCandidate]:
    while len(candidates) > MERGE_BATCH:
        batches = [
            candidates[i : i + MERGE_BATCH]
            for i in range(0, len(candidates), MERGE_BATCH)
        ]
        per_batch = max(3, math.ceil(target * 1.3 / len(batches)))
        reduced: list[ConceptCandidate] = []
        for batch in batches:
            check_cancel()
            reduced.extend(_merge_once(batch, instructions, per_batch, client))
        if len(reduced) >= len(candidates):
            # no progress; fall through to a single final merge
            candidates = reduced[:MERGE_BATCH]
            break
        candidates = reduced
    check_cancel()
    return _merge_once(candidates, instructions, target, client)
