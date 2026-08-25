# Copyright: Ankitects Pty Ltd and contributors
# License: GNU AGPL, version 3 or later; http://www.gnu.org/licenses/agpl.html

from __future__ import annotations

import os
import tempfile

from aqt.ankigpt import extract


def test_extract_docx_paragraphs_and_tables() -> None:
    import docx

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "notes.docx")
        document = docx.Document()
        document.add_heading("Supply and demand", level=1)
        document.add_paragraph("Prices move toward equilibrium.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "term"
        table.rows[0].cells[1].text = "definition"
        document.save(path)

        doc = extract.extract_text(path)
        assert doc.name == "notes.docx"
        assert "Supply and demand" in doc.text
        assert "Prices move toward equilibrium." in doc.text
        assert "term | definition" in doc.text


def test_extract_pdf() -> None:
    from pypdf import PdfWriter

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "blank.pdf")
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with open(path, "wb") as f:
            writer.write(f)
        doc = extract.extract_text(path)
        assert doc.name == "blank.pdf"
        assert doc.text == ""
